#!/usr/bin/env python3
"""Extract DOCX text, tables, comments, anchors, replies, and tracked changes."""

import argparse
import json
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
}
W_ID = f"{{{NS['w']}}}id"


def _xml(archive: zipfile.ZipFile, name: str):
    try:
        return etree.fromstring(archive.read(name))
    except KeyError:
        return None


def _text(element) -> str:
    if element is None:
        return ""
    return "".join(element.xpath(".//w:t/text() | .//w:delText/text()", namespaces=NS))


def _comments(archive: zipfile.ZipFile) -> tuple[list[dict], dict[str, str]]:
    root = _xml(archive, "word/comments.xml")
    if root is None:
        return [], {}
    comments = []
    para_to_comment: dict[str, str] = {}
    for node in root.xpath("//w:comment", namespaces=NS):
        comment_id = node.get(W_ID, "")
        para_ids = node.xpath(".//w:p/@w14:paraId", namespaces=NS)
        for para_id in para_ids:
            para_to_comment[para_id] = comment_id
        comments.append({
            "id": comment_id,
            "author": node.get(f"{{{NS['w']}}}author", ""),
            "initials": node.get(f"{{{NS['w']}}}initials", ""),
            "date": node.get(f"{{{NS['w']}}}date", ""),
            "text": _text(node),
            "anchor_text": "",
            "resolved": None,
            "parent_id": None,
        })
    return comments, para_to_comment


def _apply_extended_metadata(archive, comments: list[dict], para_to_comment: dict[str, str]) -> None:
    root = _xml(archive, "word/commentsExtended.xml")
    if root is None:
        return
    by_id = {comment["id"]: comment for comment in comments}
    for node in root.xpath("//w15:commentEx", namespaces=NS):
        para_id = node.get(f"{{{NS['w15']}}}paraId", "")
        comment_id = para_to_comment.get(para_id)
        if not comment_id or comment_id not in by_id:
            continue
        parent_para = node.get(f"{{{NS['w15']}}}paraIdParent")
        by_id[comment_id]["parent_id"] = para_to_comment.get(parent_para) if parent_para else None
        done = node.get(f"{{{NS['w15']}}}done")
        by_id[comment_id]["resolved"] = done in ("1", "true") if done is not None else None


def _anchors_and_changes(archive, comments: list[dict]) -> list[dict]:
    root = _xml(archive, "word/document.xml")
    if root is None:
        return []
    anchors: dict[str, list[str]] = {comment["id"]: [] for comment in comments}
    active: set[str] = set()
    for event, element in etree.iterwalk(root, events=("start", "end")):
        local = etree.QName(element).localname
        if event == "start" and local == "commentRangeStart":
            active.add(element.get(W_ID, ""))
        elif event == "end" and local in ("t", "delText"):
            if element.text:
                for comment_id in active:
                    anchors.setdefault(comment_id, []).append(element.text)
        elif event == "end" and local == "commentRangeEnd":
            active.discard(element.get(W_ID, ""))

    for comment in comments:
        comment["anchor_text"] = "".join(anchors.get(comment["id"], []))

    changes = []
    for kind in ("ins", "del"):
        for node in root.xpath(f"//w:{kind}", namespaces=NS):
            changes.append({
                "type": "insertion" if kind == "ins" else "deletion",
                "author": node.get(f"{{{NS['w']}}}author", ""),
                "date": node.get(f"{{{NS['w']}}}date", ""),
                "id": node.get(W_ID, ""),
                "text": _text(node),
            })
    return changes


def inspect_docx(path: Path) -> dict:
    document = Document(path)
    paragraphs = [p.text for p in document.paragraphs if p.text]
    tables = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]
    with zipfile.ZipFile(path) as archive:
        comments, para_to_comment = _comments(archive)
        _apply_extended_metadata(archive, comments, para_to_comment)
        changes = _anchors_and_changes(archive, comments)
    return {
        "file": str(path.resolve()),
        "format": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
        "comments": comments,
        "tracked_changes": changes,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    result = inspect_docx(args.input)
    encoded = json.dumps(result, ensure_ascii=False, indent=2)
    if args.output:
        args.output.write_text(encoded, encoding="utf-8")
    else:
        print(encoded)


if __name__ == "__main__":
    main()
